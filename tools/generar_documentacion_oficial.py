from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "documentacion_oficial"
LOGO = ROOT / "static" / "img" / "dulce-atardecer-logo.jpg"
NAVY, BLUE, MUTED, BORDER = "20242C", "526B85", "687385", "D6DBE2"
FOOTER = "Datanova IT Solutions · Sistema de Gestión Dulce Atardecer · Versión 1.0"


def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color); tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=NAVY, size=9):
    cell.text = ""; p = cell.paragraphs[0]; r = p.add_run(text); r.bold = bold; r.font.name = "Aptos"; r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); r.font.size = Pt(size); r.font.color.rgb = RGBColor.from_string(color); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Página ")
    field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); paragraph._p.append(field)


def setup(doc, short_title):
    sec = doc.sections[0]; sec.top_margin = Inches(.72); sec.bottom_margin = Inches(.7); sec.left_margin = Inches(.78); sec.right_margin = Inches(.78); sec.header_distance = Inches(.28); sec.footer_distance = Inches(.28)
    styles = doc.styles
    normal = styles["Normal"]; normal.font.name = "Aptos"; normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos"); normal.font.size = Pt(10); normal.font.color.rgb = RGBColor.from_string(NAVY); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [("Heading 1", 18, BLUE), ("Heading 2", 13, NAVY), ("Heading 3", 11, BLUE)]:
        style = styles[name]; style.font.name = "Aptos Display"; style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display"); style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color); style.paragraph_format.space_before = Pt(15); style.paragraph_format.space_after = Pt(7)
    hp = sec.header.paragraphs[0]; hp.text = f"Dulce Atardecer  |  {short_title}"; hp.style = styles["Normal"]; hp.runs[0].font.size = Pt(8); hp.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    fp = sec.footer.paragraphs[0]; fp.text = FOOTER + "  ·  "; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp.runs[0].font.size = Pt(7); fp.runs[0].font.color.rgb = RGBColor.from_string(MUTED); add_page_number(fp)


def cover(doc, title, subtitle, audience):
    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    if LOGO.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run().add_picture(str(LOGO), width=Inches(2.25))
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(26); r = p.add_run(title); r.bold = True; r.font.name = "Aptos Display"; r.font.size = Pt(28); r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(subtitle); r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(BLUE)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(34); r = p.add_run(audience); r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(95); r = p.add_run(f"Versión 1.0 · {date.today():%d/%m/%Y}"); r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(MUTED)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("Sistema desarrollado por\nDATANOVA IT SOLUTIONS\nwww.datanovait.com"); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_page_break()


def add_toc(doc, items):
    doc.add_heading("Índice", 1)
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="Normal"); p.paragraph_format.left_indent = Inches(.15); p.add_run(f"{i:02d}. ").bold = True; p.add_run(item)
    doc.add_page_break()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers)); table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers): shade(cell, BLUE); set_cell_text(cell, text, True, "FFFFFF", 8.5)
    for row in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row): set_cell_text(cell, str(text), False)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths): cell.width = Inches(width)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, steps):
    for item in steps:
        doc.add_paragraph(item, style="List Number")


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1); table.style = "Table Grid"; cell = table.cell(0, 0); shade(cell, "F2F4F7"); set_cell_text(cell, f"{title}: {text}", True, NAVY, 9)
    doc.add_paragraph()


def md(title, sections):
    out = [f"# {title}", "", "Sistema de Gestión Dulce Atardecer · Versión 1.0", "", "---", ""]
    for heading, body in sections:
        out += [f"## {heading}", ""]
        for block in body:
            if isinstance(block, tuple): out += [f"### {block[0]}", "", block[1], ""]
            else: out += [block, ""]
    out += ["---", "", "Sistema desarrollado por", "", "**Datanova IT Solutions**", "", "www.datanovait.com", ""]
    return "\n".join(out)


def build_delivery():
    sections = [
        ("Presentación", ["Dulce Atardecer cuenta con una aplicación web privada para centralizar la operación institucional, el seguimiento de residentes, la gestión administrativa y la coordinación diaria del personal.", "El sistema utiliza datos reales de la base institucional y está preparado para ejecutarse en PythonAnywhere con archivos estáticos administrados por WhiteNoise."]),
        ("Objetivo y alcance", ["El objetivo es reducir registros dispersos y brindar trazabilidad sobre ocupación, residentes, pagos, caja, tareas y turnos. La versión entregada no incorpora módulos de farmacia, historia clínica, liquidación de sueldos, recibos fiscales ni restauración automática de backups."]),
        ("Módulos entregados", ["Dashboard: indicadores, gráficos, filtros por geriátrico/mes/año, pagos, caja, ocupación y tareas.", "Residentes: alta web, ficha, capacidad y habitaciones, obra social, número de afiliado, cuenta corriente y estado de cuenta por email.", "Pagos y Caja: cuotas, ajustes, abonos parciales, comprobantes, ingresos automáticos, egresos, cierre y categorías.", "Personal, turnos, tareas, normas, notificaciones, perfil, configuración, reportes y PWA."]),
        ("Arquitectura y tecnologías", ["Django 5, Python 3.11, SQLite, Bootstrap 5, WhiteNoise, ReportLab, openpyxl y Chart.js. La aplicación institucional concentra modelos, formularios, vistas y reglas. Los PDFs se generan con ReportLab; los XLSX con openpyxl."]),
        ("Roles y seguridad", ["Administrador: gestión integral dentro de la aplicación y Django Admin para tareas administrativas avanzadas. Secretaría: operación diaria sin Configuración ni Administración. Consulta: lectura, exportación y sin acciones de gestión. Empleada: acceso exclusivo a sus tareas, normas, perfil, notificaciones y mis turnos.", "Las credenciales de SMTP se cargan exclusivamente como variables de entorno. SECRET_KEY, DEBUG y ALLOWED_HOSTS también se administran por entorno."]),
        ("Entrega, hosting y soporte", ["La entrega incluye código fuente, migraciones, documentación técnica, manuales, archivos estáticos y PWA. La instalación prevista es PythonAnywhere. Ejecutar migraciones, cargar datos iniciales, crear superusuario, recolectar estáticos y recargar la aplicación.", "Los accesos iniciales deben ser entregados por la persona administradora; este documento no incluye contraseñas. Datanova IT Solutions brinda soporte y actualizaciones bajo el acuerdo comercial vigente."]),
        ("Copias y actualizaciones", ["Administración puede descargar un ZIP desde Configuración con db.sqlite3, media si existe e instrucciones. La restauración es manual para evitar reemplazos accidentales. Antes de actualizar, generar un backup y ejecutar check, test y migrate en el entorno de destino."]),
    ]
    return sections


def build_admin():
    return [
        ("Antes de comenzar", ["Ingresá por /login/ con tu usuario y contraseña. Usá una cuenta Administrador para las tareas que requieren Administración Django. No compartas credenciales ni backups."]),
        ("Dashboard", ["Sirve para leer el estado institucional. Elegí geriátrico, mes y año y presioná Aplicar filtros. Revisá ocupación, pagos, caja, personal, gráficos, deuda, últimos movimientos y cumplimiento de tareas.", "Recomendación: revisar pagos vencidos y residentes con deuda antes de registrar nuevos movimientos."]),
        ("Residentes", ["Abrí Residentes y usá Nuevo residente. Completá DNI de ocho números, geriátrico, fecha de ingreso, contacto familiar, estado y datos de cobertura. La habitación debe estar disponible dentro de la capacidad del geriátrico.", "La ficha muestra cuenta corriente, total facturado, abonado y deuda. Desde allí se puede iniciar el envío del estado de cuenta. La edición detallada se realiza desde Administración."]),
        ("Pagos, cuotas y abonos", ["En Pagos se filtra por residente/DNI, geriátrico, estado y período. Registrar pago crea una cuota individual. Generar cuotas del mes crea una cuota por residente activo con monto mensual y evita duplicados por residente/período.", "Para cobrar, abrí Ver pago, completá monto, fecha, medio y observaciones en Registrar abono. El sistema permite pagos parciales, impide superar el saldo y recalcula Pendiente, Parcial, Pagado o Vencido.", "Advertencia: la fecha de vencimiento de cuotas no puede ser anterior al día actual."]),
        ("Caja", ["Los ingresos no se cargan manualmente: se generan al registrar cada abono. Para un egreso usá Registrar egreso, seleccioná geriátrico, categoría, proveedor o beneficiario, importe y medio de pago. El importe no puede superar el saldo disponible.", "Cerrar caja crea o actualiza el resumen del día con saldo inicial, ingresos, egresos, saldo final y cantidades. Usá filtros por fecha, geriátrico, categoría, proveedor y medio de pago."]),
        ("Personal, invitaciones y turnos", ["Personal lista empleados, estado laboral y estado de acceso. Una cuenta staff puede crear empleado/a y generar, copiar o regenerar invitaciones. El enlace vence a las 48 horas y se usa una sola vez.", "Turnos ofrece la grilla mensual editable M, T, N, F, L y V. Elegí mes/año, cargá códigos por empleada y guardá los cambios. M=07:00-15:00, T=15:00-23:00 y N=23:00-07:00."]),
        ("Tareas y normas", ["Las tareas y normas se administran desde Administración Django. El panel de tareas resume pendientes, completadas, vencidas y cumplimiento. Las empleadas solo ven sus propias tareas y pueden completar con observación."]),
        ("Reportes, exportaciones y emails", ["Residentes, Pagos, Caja, Personal y Turnos ofrecen exportación PDF y Excel. Los enlaces conservan filtros. Pagos permite descargar comprobante PDF y abrir una confirmación antes de enviarlo por email. Residentes permite enviar el estado de cuenta.", "Verificá destinatario, asunto y mensaje antes de Enviar. El resultado queda en el historial de envíos de Configuración."]),
        ("Configuración, usuarios y cierre", ["Configuración permite actualizar institución, logo, vencimiento/concepto de cuota, moneda y catálogos visibles. SMTP se define solo en variables de entorno. Backups requieren cuenta staff.", "Para usuarios y grupos usá Django Admin. Asigná Administrador, Secretaría o Consulta según el nivel requerido. Cerrá sesión desde el pie del menú al finalizar."]),
    ]


def build_employee():
    return [
        ("Tu acceso", ["Recibís un enlace de invitación personal. Completá usuario, email y contraseña. El enlace vence a las 48 horas y solo puede usarse una vez. Luego iniciá sesión en /login/." ]),
        ("Mis tareas", ["En Tareas se muestran únicamente las tareas asignadas a tu ficha. Revisá fecha, turno, descripción y estado. Para comenzar, presioná Iniciar. Para finalizar, elegí Marcar completada, escribí una observación si corresponde y confirmá.", "El sistema guarda automáticamente quién completó la tarea y la fecha/hora. No marques una tarea si no fue realizada."]),
        ("Mis turnos", ["En Mis turnos consultá el turno de hoy, próximo turno y semana actual. M es mañana (07:00-15:00), T tarde (15:00-23:00), N noche (23:00-07:00), F franco, L licencia y V vacaciones. Esta pantalla es solo de consulta."]),
        ("Normas y notificaciones", ["Leé las Normas y políticas activas. Cuando termines una instrucción, presioná Marcar como leído. En Notificaciones revisá tareas pendientes o vencidas y normas nuevas."]),
        ("Mi perfil y buenas prácticas", ["En Mi perfil podés actualizar nombre, apellido, email, foto opcional y contraseña. No compartas la contraseña. Cerrá sesión si usás un equipo compartido. Si una pantalla muestra un error, anotá qué estabas haciendo y comunicalo a Administración; no intentes repetir una acción financiera o administrativa."]),
    ]


def build_doc(title, subtitle, audience, sections, filename):
    doc = Document(); setup(doc, title); cover(doc, title, subtitle, audience); add_toc(doc, [h for h, _ in sections])
    for heading, blocks in sections:
        doc.add_heading(heading, 1)
        for block in blocks:
            doc.add_paragraph(block)
        if heading == "Módulos entregados":
            add_table(doc, ["Área", "Cobertura incluida"], [("Operación", "Residentes, geriátricos, habitaciones, personal y turnos"), ("Finanzas", "Cuotas, abonos, estados, caja, cierres y exportaciones"), ("Personas", "Tareas, normas, invitaciones, perfil y notificaciones"), ("Plataforma", "Roles, PWA, backups, archivos estáticos y reportes")])
        if heading == "Roles y seguridad":
            add_note(doc, "Importante", "Las cuentas de Consulta no pueden ejecutar acciones de gestión aunque intenten acceder por URL directa.")
        if heading == "Pagos, cuotas y abonos":
            add_steps(doc, ["Abrí Pagos y localizá o registrá la cuota.", "Ingresá al detalle con Ver pago.", "Completá el abono sin superar el saldo pendiente.", "Verificá el estado y el movimiento automático en Caja."])
        if heading == "Mis tareas":
            add_steps(doc, ["Ingresá a Tareas.", "Leé el detalle y verificá el turno.", "Presioná Iniciar cuando comiences.", "Al terminar, elegí Marcar completada y confirmá la observación."])
    doc.add_page_break(); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(110); r = p.add_run("Sistema desarrollado por\n\nDatanova IT Solutions\nwww.datanovait.com"); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(NAVY)
    path = OUT / f"{filename}.docx"; doc.save(path)
    (OUT / f"{filename}.md").write_text(md(title, sections), encoding="utf-8")
    build_pdf(title, subtitle, audience, sections, filename)


def build_pdf(title, subtitle, audience, sections, filename):
    path = OUT / f"{filename}.pdf"
    pdf = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=1.7*cm, leftMargin=1.7*cm, topMargin=1.9*cm, bottomMargin=1.8*cm)
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.4, leading=13, textColor=colors.HexColor("#20242C"), spaceAfter=7)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=17, leading=21, textColor=colors.HexColor("#526B85"), spaceBefore=14, spaceAfter=8)
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=26, leading=31, alignment=1, textColor=colors.HexColor("#20242C"), spaceAfter=8)
    sub = ParagraphStyle("Sub", parent=body, fontSize=13, leading=17, alignment=1, textColor=colors.HexColor("#526B85"), spaceAfter=18)
    center = ParagraphStyle("Center", parent=body, alignment=1, textColor=colors.HexColor("#687385"))
    story = [Spacer(1, 2.2*cm)]
    if LOGO.exists(): story += [Image(str(LOGO), width=5.2*cm, height=5.2*cm), Spacer(1, .5*cm)]
    story += [Paragraph(title, title_style), Paragraph(subtitle, sub), Paragraph(audience, center), Spacer(1, 3.0*cm), Paragraph(f"Versión 1.0 · {date.today():%d/%m/%Y}", center), Spacer(1, .5*cm), Paragraph("Sistema desarrollado por<br/><b>DATANOVA IT SOLUTIONS</b><br/>www.datanovait.com", center), PageBreak(), Paragraph("Índice", h1)]
    for i, (heading, _) in enumerate(sections, 1): story.append(Paragraph(f"{i:02d}. {heading}", body))
    story.append(PageBreak())
    for heading, blocks in sections:
        story.append(Paragraph(heading, h1))
        for block in blocks: story.append(Paragraph(block.replace("&", "&amp;"), body))
        if heading == "Módulos entregados":
            data = [["Área", "Cobertura incluida"], ["Operación", "Residentes, geriátricos, habitaciones, personal y turnos"], ["Finanzas", "Cuotas, abonos, caja, cierres y exportaciones"], ["Personas", "Tareas, normas, invitaciones, perfil y notificaciones"], ["Plataforma", "Roles, PWA, backups, estáticos y reportes"]]
            table = Table(data, colWidths=[3.2*cm, 12.7*cm], repeatRows=1)
            table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.HexColor("#526B85")), ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"), ("FONTSIZE", (0,0), (-1,-1), 8.5), ("GRID", (0,0), (-1,-1), .25, colors.HexColor("#D6DBE2")), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("TOPPADDING", (0,0), (-1,-1), 6), ("BOTTOMPADDING", (0,0), (-1,-1), 6)])); story += [table, Spacer(1, .35*cm)]
        if heading == "Pagos, cuotas y abonos":
            story += [Paragraph("Secuencia recomendada", h1), Paragraph("1. Abrí Pagos y localizá o registrá la cuota.<br/>2. Ingresá al detalle con Ver pago.<br/>3. Completá el abono sin superar el saldo pendiente.<br/>4. Verificá el estado y el ingreso automático en Caja.", body)]
        if heading == "Mis tareas":
            story += [Paragraph("Pasos de trabajo", h1), Paragraph("1. Ingresá a Tareas.<br/>2. Leé el detalle y verificá el turno.<br/>3. Presioná Iniciar cuando comiences.<br/>4. Al terminar, elegí Marcar completada y confirmá la observación.", body)]
    story += [PageBreak(), Spacer(1, 8*cm), Paragraph("Sistema desarrollado por<br/><br/><b>Datanova IT Solutions</b><br/>www.datanovait.com", center)]
    def footer(canvas, doc):
        canvas.saveState(); canvas.setStrokeColor(colors.HexColor("#D6DBE2")); canvas.line(doc.leftMargin, 1.25*cm, A4[0]-doc.rightMargin, 1.25*cm); canvas.setFillColor(colors.HexColor("#687385")); canvas.setFont("Helvetica", 7); canvas.drawCentredString(A4[0]/2, .83*cm, FOOTER); canvas.drawRightString(A4[0]-doc.rightMargin, .56*cm, f"Página {canvas.getPageNumber()}"); canvas.restoreState()
    pdf.build(story, onFirstPage=footer, onLaterPages=footer)


def main():
    OUT.mkdir(exist_ok=True)
    build_doc("Documento de Entrega", "Sistema de Gestión para Geriátrico Dulce Atardecer", "Documento de entrega al cliente", build_delivery(), "documento_entrega")
    build_doc("Manual de Usuario", "Administración y dueños", "Guía operativa para gestión administrativa", build_admin(), "manual_usuario_administracion")
    build_doc("Manual para Empleadas", "Uso diario del sistema", "Guía práctica para el personal del geriátrico", build_employee(), "manual_empleadas")


if __name__ == "__main__": main()
